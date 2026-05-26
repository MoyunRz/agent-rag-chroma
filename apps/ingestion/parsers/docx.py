from .base import BaseParser, ParsedDocument, ParsedSection


class DocxParser(BaseParser):
    """Word 文档解析器 (python-docx)."""

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError('Word 解析需要安装 python-docx: pip install python-docx')

        doc = DocxDocument(file_path)
        sections = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            level = para.style.name if para.style else ''
            is_heading = level.startswith('Heading')
            heading_level = int(level[-1]) if is_heading and level[-1].isdigit() else 0
            sections.append(ParsedSection(
                text=text,
                heading_level=heading_level,
                heading=text if is_heading else '',
            ))
        title = file_path.rsplit('/', 1)[-1].rsplit('\\', 1)[-1]
        raw = '\n'.join(s.text for s in sections)
        return ParsedDocument(title=title, sections=sections, raw_text=raw)
